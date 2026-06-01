# -*- coding: utf-8 -*-
"""
根据已编码的外交部例行记者会 Excel 文件生成统计表

输入：
mfa_press_2026_03_04_.xlsx

输出：
1. mfa_press_统计表.xlsx
2. mfa_press_统计表.docx
"""

import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# =========================
# 1. 文件路径设置
# =========================

INPUT_FILE = "/Users/lvqinyang/Desktop/mfa_press_2026_03_04_.xlsx"

OUTPUT_EXCEL = "mfa_press_统计表.xlsx"
OUTPUT_WORD = "mfa_press_统计表.docx"

SHEET_NAME = "qa_units"


# =========================
# 2. 读取数据
# =========================

df = pd.read_excel(INPUT_FILE, sheet_name=SHEET_NAME)

# 删除没有有效回答的行
df = df[df["answer"].notna()].copy()

# 如果 issue_type 为空，标记为“未编码”
df["issue_type"] = df["issue_type"].fillna("未编码")

# 话语策略变量
strategy_vars = [
    "cooperation_frame",
    "development_frame",
    "sovereignty_frame",
    "multilateral_frame",
    "responsibility_attribution",
    "criticism_frame",
    "positive_self_image",
]

strategy_names = {
    "cooperation_frame": "合作框架",
    "development_frame": "发展框架",
    "sovereignty_frame": "主权安全框架",
    "multilateral_frame": "多边主义框架",
    "responsibility_attribution": "责任归因",
    "criticism_frame": "批评性话语",
    "positive_self_image": "中国正面形象建构",
}

# 将策略变量统一转成 0/1
for col in strategy_vars:
    if col not in df.columns:
        df[col] = 0
    df[col] = pd.to_numeric(df[col], errors="coerce").fillna(0).astype(int)


# =========================
# 3. 表1：样本基本情况
# =========================

sample_start = df["date"].min()
sample_end = df["date"].max()

press_count = df["date"].nunique()
qa_count = len(df)

if "title" in df.columns:
    title_count = df["title"].nunique()
else:
    title_count = press_count

table_sample = pd.DataFrame({
    "项目": [
        "样本时间范围",
        "记者会场次",
        "有效问答单元数",
        "数据来源",
        "分析单位",
    ],
    "数值": [
        f"{sample_start} 至 {sample_end}",
        f"{press_count} 场",
        f"{qa_count} 个",
        "中华人民共和国外交部官网例行记者会栏目",
        "记者提问—发言人回答",
    ]
})


# =========================
# 4. 表2：议题类型分布
# =========================

issue_order = [
    "大国关系",
    "周边外交",
    "主权安全",
    "全球治理",
    "经贸科技",
    "国际热点",
    "人文交流",
    "其他",
    "未编码",
]

issue_counts = df["issue_type"].value_counts()

table_issue = (
    issue_counts
    .rename_axis("议题类型")
    .reset_index(name="频数")
)

# 按预设顺序排序
table_issue["排序"] = table_issue["议题类型"].apply(
    lambda x: issue_order.index(x) if x in issue_order else 999
)
table_issue = table_issue.sort_values("排序").drop(columns="排序")

table_issue["占比"] = table_issue["频数"] / qa_count
table_issue["占比"] = table_issue["占比"].apply(lambda x: f"{x:.2%}")


# =========================
# 5. 表3：话语策略分布
# =========================

strategy_rows = []

for col in strategy_vars:
    count = int(df[col].sum())
    ratio = count / qa_count if qa_count > 0 else 0

    strategy_rows.append({
        "话语策略": strategy_names[col],
        "出现次数": count,
        "占比": f"{ratio:.2%}",
    })

table_strategy = pd.DataFrame(strategy_rows)


# =========================
# 6. 附加表：议题类型 × 话语策略交叉表
# =========================

cross_rows = []

for issue_type, group in df.groupby("issue_type"):
    row = {
        "议题类型": issue_type,
        "样本数": len(group),
    }

    for col in strategy_vars:
        row[strategy_names[col]] = int(group[col].sum())

    cross_rows.append(row)

table_cross = pd.DataFrame(cross_rows)

# 按议题顺序排序
table_cross["排序"] = table_cross["议题类型"].apply(
    lambda x: issue_order.index(x) if x in issue_order else 999
)
table_cross = table_cross.sort_values("排序").drop(columns="排序")


# =========================
# 7. 生成 Excel 文件
# =========================

with pd.ExcelWriter(OUTPUT_EXCEL, engine="openpyxl") as writer:
    table_sample.to_excel(writer, sheet_name="表1_样本基本情况", index=False)
    table_issue.to_excel(writer, sheet_name="表2_议题类型分布", index=False)
    table_strategy.to_excel(writer, sheet_name="表3_话语策略分布", index=False)
    table_cross.to_excel(writer, sheet_name="附表_议题策略交叉表", index=False)

    workbook = writer.book

    for sheet_name in workbook.sheetnames:
        ws = workbook[sheet_name]

        # 冻结首行
        ws.freeze_panes = "A2"

        # 自动筛选
        ws.auto_filter.ref = ws.dimensions

        # 设置表头样式
        for cell in ws[1]:
            cell.font = cell.font.copy(bold=True)

        # 设置列宽
        for col in ws.columns:
            max_length = 0
            col_letter = col[0].column_letter

            for cell in col:
                value = str(cell.value) if cell.value is not None else ""
                max_length = max(max_length, len(value))

            adjusted_width = min(max_length + 4, 35)
            ws.column_dimensions[col_letter].width = adjusted_width


# =========================
# 8. 生成 Word 文件
# =========================

def add_table_to_doc(doc, title, dataframe):
    """把 DataFrame 添加到 Word 文档中。"""
    doc.add_heading(title, level=2)

    table = doc.add_table(rows=1, cols=len(dataframe.columns))
    table.style = "Table Grid"

    # 表头
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(dataframe.columns):
        hdr_cells[i].text = str(col_name)

    # 内容
    for _, row in dataframe.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

    doc.add_paragraph("")


doc = Document()

title = doc.add_heading("外交部例行记者会文本内容分析统计表", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph(
    "本文件根据已编码的外交部例行记者会问答单元生成，主要包括样本基本情况、议题类型分布、话语策略分布以及议题类型与话语策略交叉表。"
)
p.paragraph_format.first_line_indent = Pt(21)

add_table_to_doc(doc, "表1 样本基本情况", table_sample)
add_table_to_doc(doc, "表2 议题类型分布", table_issue)
add_table_to_doc(doc, "表3 话语策略分布", table_strategy)
add_table_to_doc(doc, "附表 议题类型与话语策略交叉表", table_cross)

doc.save(OUTPUT_WORD)


# =========================
# 9. 输出提示
# =========================

print("统计表生成完成！")
print(f"Excel 文件：{OUTPUT_EXCEL}")
print(f"Word 文件：{OUTPUT_WORD}")

print("\n表1：样本基本情况")
print(table_sample)

print("\n表2：议题类型分布")
print(table_issue)

print("\n表3：话语策略分布")
print(table_strategy)

print("\n附表：议题类型 × 话语策略交叉表")
print(table_cross)
