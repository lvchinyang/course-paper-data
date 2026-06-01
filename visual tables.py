# -*- coding: utf-8 -*-
"""
根据外交部例行记者会内容分析结果生成高级可视化图表

输入文件：
优先读取 mfa_press_统计表.xlsx
如果没有这个文件，也可以改成读取 mfa_press_2026_03_04_已编码.xlsx 后自行统计。

输出：
1. chart_issue_distribution.png
2. chart_strategy_lollipop.png
3. chart_issue_strategy_heatmap.png
4. chart_issue_strategy_bubble.png
5. 外交部例行记者会_可视化图表.docx

运行前安装：
pip install pandas openpyxl matplotlib python-docx
"""

import os
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# =========================
# 1. 文件路径设置
# =========================

INPUT_FILE = "mfa_press_统计表.xlsx"

OUT_DIR = "mfa_visual_outputs"
os.makedirs(OUT_DIR, exist_ok=True)

OUT_ISSUE_BAR = os.path.join(OUT_DIR, "chart_issue_distribution.png")
OUT_STRATEGY_LOLLIPOP = os.path.join(OUT_DIR, "chart_strategy_lollipop.png")
OUT_HEATMAP = os.path.join(OUT_DIR, "chart_issue_strategy_heatmap.png")
OUT_BUBBLE = os.path.join(OUT_DIR, "chart_issue_strategy_bubble.png")
OUT_WORD = os.path.join(OUT_DIR, "外交部例行记者会_可视化图表.docx")


# =========================
# 2. 中文字体设置
# =========================
# Mac 常见中文字体：Arial Unicode MS、PingFang SC、Songti SC
# Windows 可以改成：SimHei 或 Microsoft YaHei

matplotlib.rcParams["font.sans-serif"] = [
    "Arial Unicode MS",
    "PingFang SC",
    "Songti SC",
    "SimHei",
    "Microsoft YaHei",
]
matplotlib.rcParams["axes.unicode_minus"] = False


# =========================
# 3. 读取统计表
# =========================

issue_df = pd.read_excel(INPUT_FILE, sheet_name="表2_议题类型分布")
strategy_df = pd.read_excel(INPUT_FILE, sheet_name="表3_话语策略分布")
cross_df = pd.read_excel(INPUT_FILE, sheet_name="附表_议题策略交叉表")

# 兼容占比为字符串百分比的情况
def percent_to_float(x):
    if isinstance(x, str):
        return float(x.replace("%", "")) / 100
    return float(x)

issue_df["占比_float"] = issue_df["占比"].apply(percent_to_float)
strategy_df["占比_float"] = strategy_df["占比"].apply(percent_to_float)

# 议题顺序
issue_order = [
    "国际热点",
    "主权安全",
    "经贸科技",
    "周边外交",
    "大国关系",
    "人文交流",
    "全球治理",
    "其他",
]

issue_df["排序"] = issue_df["议题类型"].apply(
    lambda x: issue_order.index(x) if x in issue_order else 999
)
issue_df = issue_df.sort_values("排序")

cross_df["排序"] = cross_df["议题类型"].apply(
    lambda x: issue_order.index(x) if x in issue_order else 999
)
cross_df = cross_df.sort_values("排序").drop(columns=["排序"])


# =========================
# 4. 图1：议题类型分布横向条形图
# =========================

def plot_issue_distribution():
    df = issue_df.sort_values("频数", ascending=True)

    fig, ax = plt.subplots(figsize=(9, 5.5))
    bars = ax.barh(df["议题类型"], df["频数"])

    ax.set_title("图1 议题类型分布", fontsize=16, pad=14)
    ax.set_xlabel("问答单元数量", fontsize=11)
    ax.set_ylabel("议题类型", fontsize=11)

    # 添加数值和占比标签
    for bar, freq, ratio in zip(bars, df["频数"], df["占比_float"]):
        width = bar.get_width()
        ax.text(
            width + max(df["频数"]) * 0.015,
            bar.get_y() + bar.get_height() / 2,
            f"{freq}（{ratio:.2%}）",
            va="center",
            fontsize=10,
        )

    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="x", linestyle="--", alpha=0.35)

    plt.tight_layout()
    plt.savefig(OUT_ISSUE_BAR, dpi=300, bbox_inches="tight")
    plt.close()


# =========================
# 5. 图2：话语策略棒棒糖图
# =========================

def plot_strategy_lollipop():
    df = strategy_df.sort_values("出现次数", ascending=True)

    fig, ax = plt.subplots(figsize=(9, 5.5))

    y_pos = range(len(df))
    ax.hlines(
        y=y_pos,
        xmin=0,
        xmax=df["出现次数"],
        linewidth=2,
        alpha=0.75,
    )
    ax.scatter(
        df["出现次数"],
        y_pos,
        s=130,
        zorder=3,
    )

    ax.set_yticks(y_pos)
    ax.set_yticklabels(df["话语策略"])
    ax.set_xlabel("出现次数", fontsize=11)
    ax.set_title("图2 话语策略出现频率", fontsize=16, pad=14)

    for x, y, ratio in zip(df["出现次数"], y_pos, df["占比_float"]):
        ax.text(
            x + max(df["出现次数"]) * 0.02,
            y,
            f"{int(x)}（{ratio:.2%}）",
            va="center",
            fontsize=10,
        )

    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="x", linestyle="--", alpha=0.35)

    plt.tight_layout()
    plt.savefig(OUT_STRATEGY_LOLLIPOP, dpi=300, bbox_inches="tight")
    plt.close()


# =========================
# 6. 图3：议题类型 × 话语策略热力图
# =========================

def plot_heatmap():
    strategy_cols = [
        "合作框架",
        "发展框架",
        "主权安全框架",
        "多边主义框架",
        "责任归因",
        "批评性话语",
        "中国正面形象建构",
    ]

    matrix = cross_df.set_index("议题类型")[strategy_cols]

    fig, ax = plt.subplots(figsize=(10.5, 6))

    im = ax.imshow(matrix.values, aspect="auto")

    ax.set_xticks(range(len(strategy_cols)))
    ax.set_xticklabels(strategy_cols, rotation=35, ha="right")
    ax.set_yticks(range(len(matrix.index)))
    ax.set_yticklabels(matrix.index)

    ax.set_title("图3 议题类型与话语策略交叉热力图", fontsize=16, pad=14)

    # 添加数值标签
    for i in range(matrix.shape[0]):
        for j in range(matrix.shape[1]):
            value = matrix.iloc[i, j]
            ax.text(
                j,
                i,
                str(value),
                ha="center",
                va="center",
                fontsize=9,
            )

    cbar = plt.colorbar(im, ax=ax)
    cbar.set_label("出现次数", fontsize=10)

    plt.tight_layout()
    plt.savefig(OUT_HEATMAP, dpi=300, bbox_inches="tight")
    plt.close()


# =========================
# 7. 图4：议题类型 × 话语策略气泡矩阵
# =========================

def plot_bubble_matrix():
    strategy_cols = [
        "合作框架",
        "发展框架",
        "主权安全框架",
        "多边主义框架",
        "责任归因",
        "批评性话语",
        "中国正面形象建构",
    ]

    matrix = cross_df.set_index("议题类型")[strategy_cols]

    x_labels = strategy_cols
    y_labels = list(matrix.index)

    xs, ys, sizes, values = [], [], [], []

    max_value = matrix.values.max()

    for i, issue in enumerate(y_labels):
        for j, strategy in enumerate(x_labels):
            value = matrix.loc[issue, strategy]
            xs.append(j)
            ys.append(i)
            values.append(value)

            # 气泡大小按次数缩放
            size = 80 + (value / max_value) * 900 if max_value > 0 else 80
            sizes.append(size)

    fig, ax = plt.subplots(figsize=(11, 6.2))

    scatter = ax.scatter(
        xs,
        ys,
        s=sizes,
        c=values,
        alpha=0.75,
    )

    ax.set_xticks(range(len(x_labels)))
    ax.set_xticklabels(x_labels, rotation=35, ha="right")
    ax.set_yticks(range(len(y_labels)))
    ax.set_yticklabels(y_labels)

    ax.set_title("图4 议题类型与话语策略气泡矩阵", fontsize=16, pad=14)

    # 添加数值
    for x, y, value in zip(xs, ys, values):
        if value > 0:
            ax.text(
                x,
                y,
                str(value),
                ha="center",
                va="center",
                fontsize=8,
            )

    cbar = plt.colorbar(scatter, ax=ax)
    cbar.set_label("出现次数", fontsize=10)

    ax.grid(True, linestyle="--", alpha=0.25)

    plt.tight_layout()
    plt.savefig(OUT_BUBBLE, dpi=300, bbox_inches="tight")
    plt.close()


# =========================
# 8. 生成 Word 文档
# =========================

def add_image_with_caption(doc, image_path, caption, width=6.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width))

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def make_word_report():
    doc = Document()

    title = doc.add_heading("外交部例行记者会文本内容分析可视化图表", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "本文件根据已编码的外交部例行记者会问答单元生成，主要展示议题类型分布、话语策略分布以及议题类型与话语策略之间的交叉关系。"
    )

    add_image_with_caption(
        doc,
        OUT_ISSUE_BAR,
        "图1 议题类型分布。国际热点、主权安全和经贸科技是样本期内较突出的议题类型。",
        width=6.3,
    )

    doc.add_page_break()

    add_image_with_caption(
        doc,
        OUT_STRATEGY_LOLLIPOP,
        "图2 话语策略出现频率。合作框架和责任归因是出现频率最高的两类话语策略。",
        width=6.3,
    )

    doc.add_page_break()

    add_image_with_caption(
        doc,
        OUT_HEATMAP,
        "图3 议题类型与话语策略交叉热力图。颜色越深表示对应话语策略在该议题类型中出现越多。",
        width=6.5,
    )

    doc.add_page_break()

    add_image_with_caption(
        doc,
        OUT_BUBBLE,
        "图4 议题类型与话语策略气泡矩阵。气泡越大表示对应组合出现次数越高。",
        width=6.5,
    )

    doc.save(OUT_WORD)


# =========================
# 9. 主程序
# =========================

def main():
    print("开始生成可视化图表……")

    plot_issue_distribution()
    print(f"已生成：{OUT_ISSUE_BAR}")

    plot_strategy_lollipop()
    print(f"已生成：{OUT_STRATEGY_LOLLIPOP}")

    plot_heatmap()
    print(f"已生成：{OUT_HEATMAP}")

    plot_bubble_matrix()
    print(f"已生成：{OUT_BUBBLE}")

    make_word_report()
    print(f"已生成：{OUT_WORD}")

    print("\n全部完成！")
    print(f"图表文件夹：{OUT_DIR}")


if __name__ == "__main__":
    main()