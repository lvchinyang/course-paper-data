# -*- coding: utf-8 -*-
"""
抓取外交部官网 2026年3月—4月 例行记者会文本

输出文件：
1. mfa_press_2026_03_04.txt   按时间顺序整理的纯文本
2. mfa_press_2026_03_04.docx  Word文档
3. mfa_press_2026_03_04.xlsx  按“记者提问—发言人回答”切分的Excel编码表

运行前安装：
pip install requests beautifulsoup4 lxml pandas python-docx openpyxl
"""

import re
import time
from datetime import datetime
from urllib.parse import urljoin

import requests
import pandas as pd
from bs4 import BeautifulSoup
from docx import Document
from openpyxl.styles import Alignment, Font, PatternFill


# =========================================================
# 1. 基础配置
# =========================================================

BASE_URL = "https://www.mfa.gov.cn/web/wjdt_674879/fyrbt_674889/"

START_DATE = datetime(2026, 3, 1)
END_DATE = datetime(2026, 4, 30)

OUT_TXT = "mfa_press_2026_03_04.txt"
OUT_DOCX = "mfa_press_2026_03_04.docx"
OUT_XLSX = "mfa_press_2026_03_04.xlsx"

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0 Safari/537.36"
    )
}


# =========================================================
# 2. 通用工具函数
# =========================================================

def get_html(url: str, timeout: int = 20) -> str:
    """请求网页 HTML。"""
    resp = requests.get(url, headers=HEADERS, timeout=timeout)
    resp.encoding = resp.apparent_encoding or "utf-8"
    resp.raise_for_status()
    return resp.text


def clean_text(text: str) -> str:
    """清理空白字符。"""
    text = text.replace("\u3000", " ")
    text = re.sub(r"\r\n|\r", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def parse_date_from_title(title: str):
    """
    从标题中提取日期。
    示例：
    2026年4月30日外交部发言人林剑主持例行记者会
    """
    m = re.search(r"(\d{4})年(\d{1,2})月(\d{1,2})日", title)
    if not m:
        m = re.search(r"(\d{4})-(\d{1,2})-(\d{1,2})", title)

    if not m:
        return None

    y, mo, d = map(int, m.groups())
    return datetime(y, mo, d)


def looks_like_press_conference(title: str) -> bool:
    """判断标题是否像外交部例行记者会。"""
    return "例行记者会" in title and "外交部发言人" in title


# =========================================================
# 3. 收集栏目页和文章链接
# =========================================================

def collect_list_page_urls(max_pages: int = 40):
    """
    自动构造外交部例行记者会栏目分页 URL。

    常见形式：
    - /fyrbt_674889/
    - /fyrbt_674889/index.shtml
    - /fyrbt_674889/index_1.shtml
    - /fyrbt_674889/index_2.shtml
    """
    candidates = [
        BASE_URL,
        urljoin(BASE_URL, "index.shtml"),
    ]

    for i in range(1, max_pages + 1):
        candidates.append(urljoin(BASE_URL, f"index_{i}.shtml"))
        candidates.append(urljoin(BASE_URL, f"index_{i}.htm"))
        candidates.append(urljoin(BASE_URL, f"index_{i}.html"))

    seen = set()
    urls = []
    for u in candidates:
        if u not in seen:
            urls.append(u)
            seen.add(u)

    return urls


def collect_article_links():
    """
    从栏目页收集 2026年3—4月例行记者会文章链接。
    """
    article_map = {}

    print("开始检查栏目分页……")

    for list_url in collect_list_page_urls(max_pages=40):
        try:
            html = get_html(list_url)
        except Exception:
            continue

        soup = BeautifulSoup(html, "lxml")
        links = soup.find_all("a")

        found_on_this_page = 0

        for a in links:
            title = a.get_text(strip=True)
            href = a.get("href")

            if not title or not href:
                continue

            if not looks_like_press_conference(title):
                continue

            date_obj = parse_date_from_title(title)
            if not date_obj:
                continue

            if START_DATE <= date_obj <= END_DATE:
                article_url = urljoin(list_url, href)

                article_map[article_url] = {
                    "date": date_obj,
                    "title": title,
                    "url": article_url,
                }

                found_on_this_page += 1

        if found_on_this_page > 0:
            print(f"栏目页：{list_url}，找到目标链接 {found_on_this_page} 条")

        time.sleep(0.4)

    articles = list(article_map.values())
    articles.sort(key=lambda x: x["date"])

    return articles


# =========================================================
# 4. 抓取正文
# =========================================================

def extract_article_text(url: str):
    """
    提取记者会正文。
    适配外交部网页常见正文结构。
    """
    html = get_html(url)
    soup = BeautifulSoup(html, "lxml")

    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    selectors = [
        "div.TRS_Editor",
        "div.trs_editor",
        "div#News_Body_Txt_A",
        "div.news_content",
        "div.content",
        "article",
    ]

    content_node = None

    for sel in selectors:
        node = soup.select_one(sel)
        if node and len(node.get_text(strip=True)) > 200:
            content_node = node
            break

    if content_node is None:
        divs = soup.find_all("div")
        if divs:
            content_node = max(divs, key=lambda d: len(d.get_text(strip=True)))
        else:
            content_node = soup.body or soup

    paragraphs = []
    for p in content_node.find_all(["p"], recursive=True):
        txt = p.get_text(" ", strip=True)
        if txt and len(txt) > 2:
            paragraphs.append(txt)

    if len(paragraphs) < 5:
        paragraphs = []
        for p in content_node.find_all(["p", "div"], recursive=True):
            txt = p.get_text(" ", strip=True)
            if txt and len(txt) > 2:
                paragraphs.append(txt)

    if not paragraphs:
        text = content_node.get_text("\n", strip=True)
        return clean_text(text)

    seen = set()
    unique_paragraphs = []

    for para in paragraphs:
        para = clean_text(para)
        if para and para not in seen:
            unique_paragraphs.append(para)
            seen.add(para)

    text = "\n\n".join(unique_paragraphs)
    return clean_text(text)


# =========================================================
# 5. 增强版问答切分
# =========================================================

def split_qa_units(full_text: str):
    """
    段落级增强版问答切分函数。

    目标：
    把外交部例行记者会正文切分成多个“记者提问—发言人回答”单元。
    """

    text = full_text.replace("\u3000", " ")
    text = re.sub(r"\r\n|\r", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.strip()

    if "\n\n" in text:
        parts = [p.strip() for p in text.split("\n\n") if p.strip()]
    else:
        parts = [p.strip() for p in text.split("\n") if p.strip()]

    spokesperson_names = [
        "毛宁",
        "林剑",
        "郭嘉昆",
        "汪文斌",
        "华春莹",
        "赵立坚",
        "陆慷",
        "耿爽",
        "秦刚",
    ]

    question_patterns = [
        r"^问[：:]",
        r"^记者[：:]",
        r"^追问[：:]",
        r"^补充提问[：:]",
        r"^.+?记者[：:]",
        r"^.+?社记者[：:]",
        r"^.+?电视台记者[：:]",
        r"^.+?广播公司记者[：:]",
        r"^.+?报社记者[：:]",
        r"^.+?时报记者[：:]",
        r"^.+?新闻记者[：:]",
        r"^.+?通讯社记者[：:]",
        r"^.+?日报记者[：:]",
        r"^.+?卫视记者[：:]",
        r"^.+?电台记者[：:]",
        r"^.+?杂志记者[：:]",
    ]

    answer_patterns = [
        r"^答[：:]",
        r"^发言人[：:]",
    ]

    for name in spokesperson_names:
        answer_patterns.append(rf"^{re.escape(name)}[：:]")

    def is_question_start(s: str) -> bool:
        return any(re.match(p, s) for p in question_patterns)

    def is_answer_start(s: str) -> bool:
        return any(re.match(p, s) for p in answer_patterns)

    qa_units = []
    current_q = []
    current_a = []
    mode = None

    def save_current():
        q = "\n".join(current_q).strip()
        a = "\n".join(current_a).strip()

        if q or a:
            qa_units.append({
                "question": q,
                "answer": a
            })

    for part in parts:
        part = part.strip()
        if not part:
            continue

        if is_question_start(part):
            if current_q or current_a:
                save_current()

            current_q = [part]
            current_a = []
            mode = "q"

        elif is_answer_start(part):
            current_a.append(part)
            mode = "a"

        else:
            if mode == "q":
                current_q.append(part)
            elif mode == "a":
                current_a.append(part)
            else:
                continue

    if current_q or current_a:
        save_current()

    cleaned = []

    for unit in qa_units:
        q = unit["question"].strip()
        a = unit["answer"].strip()

        if not a:
            for name in spokesperson_names:
                pattern = rf"({re.escape(name)}[：:].*)"
                m = re.search(pattern, q)
                if m:
                    q_part = q[:m.start()].strip()
                    a_part = q[m.start():].strip()

                    q = q_part
                    a = a_part
                    break

        if not a:
            m = re.search(r"(答[：:].*)", q)
            if m:
                q_part = q[:m.start()].strip()
                a_part = q[m.start():].strip()
                q = q_part
                a = a_part

        if len(q) >= 5 or len(a) >= 5:
            cleaned.append({
                "question": q,
                "answer": a
            })

    return cleaned


# =========================================================
# 6. 保存 TXT / DOCX / XLSX
# =========================================================

def save_txt(records):
    """保存完整文本为 TXT，并在每一场记者会前显示原文链接。"""
    with open(OUT_TXT, "w", encoding="utf-8") as f:
        f.write("2026年3月—4月外交部例行记者会文本汇总\n")
        f.write("=" * 70 + "\n\n")

        for rec in records:
            f.write(f"日期：{rec['date_str']}\n")
            f.write(f"标题：{rec['title']}\n")
            f.write(f"原文链接：{rec['url']}\n")
            f.write("-" * 70 + "\n\n")
            f.write(rec["text"])
            f.write("\n\n" + "=" * 70 + "\n\n")


def save_docx(records):
    """保存完整文本为 Word，并在每一场记者会标题下显示原文链接。"""
    doc = Document()
    doc.add_heading("2026年3月—4月外交部例行记者会文本汇总", level=0)

    for rec in records:
        doc.add_heading(f"{rec['date_str']}  {rec['title']}", level=1)

        p1 = doc.add_paragraph()
        p1.add_run("日期：").bold = True
        p1.add_run(rec["date_str"])

        p2 = doc.add_paragraph()
        p2.add_run("原文链接：").bold = True
        p2.add_run(rec["url"])

        doc.add_paragraph("")

        for para in rec["text"].split("\n\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)

        doc.add_page_break()

    doc.save(OUT_DOCX)


def save_excel(records):
    """
    保存问答单元 Excel。
    每个问答单元都保留日期、标题、原文链接。
    增加拆分质量检查列，方便后续人工复核和内容分析编码。
    """
    rows = []

    for rec in records:
        qa_units = split_qa_units(rec["text"])

        if not qa_units:
            rows.append({
                "date": rec["date_str"],
                "title": rec["title"],
                "source_url": rec["url"],
                "qa_id": f"{rec['date_str']}-00",
                "question": "",
                "answer": rec["text"],
                "split_status": "未成功拆分",
                "question_len": 0,
                "answer_len": len(rec["text"]),
                "needs_check": 1,

                "issue_type": "",
                "cooperation_frame": "",
                "development_frame": "",
                "sovereignty_frame": "",
                "multilateral_frame": "",
                "responsibility_attribution": "",
                "criticism_frame": "",
                "positive_self_image": "",
                "notes": "",
            })

        else:
            for i, qa in enumerate(qa_units, start=1):
                q = qa.get("question", "").strip()
                a = qa.get("answer", "").strip()

                needs_check = 0
                split_status = "成功拆分"

                if not q or not a:
                    needs_check = 1
                    split_status = "问题或回答为空"
                elif len(q) < 10:
                    needs_check = 1
                    split_status = "问题过短"
                elif len(a) < 20:
                    needs_check = 1
                    split_status = "回答过短"

                rows.append({
                    "date": rec["date_str"],
                    "title": rec["title"],
                    "source_url": rec["url"],
                    "qa_id": f"{rec['date_str']}-{i:02d}",
                    "question": q,
                    "answer": a,
                    "split_status": split_status,
                    "question_len": len(q),
                    "answer_len": len(a),
                    "needs_check": needs_check,

                    "issue_type": "",
                    "cooperation_frame": "",
                    "development_frame": "",
                    "sovereignty_frame": "",
                    "multilateral_frame": "",
                    "responsibility_attribution": "",
                    "criticism_frame": "",
                    "positive_self_image": "",
                    "notes": "",
                })

    df = pd.DataFrame(rows)

    if not df.empty:
        df = df.sort_values(by=["date", "qa_id"])

    with pd.ExcelWriter(OUT_XLSX, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name="qa_units", index=False)

        worksheet = writer.sheets["qa_units"]

        worksheet.freeze_panes = "A2"
        worksheet.auto_filter.ref = worksheet.dimensions

        column_widths = {
            "A": 12,   # date
            "B": 42,   # title
            "C": 65,   # source_url
            "D": 18,   # qa_id
            "E": 55,   # question
            "F": 70,   # answer
            "G": 18,   # split_status
            "H": 12,   # question_len
            "I": 12,   # answer_len
            "J": 12,   # needs_check
            "K": 18,   # issue_type
            "L": 20,   # cooperation_frame
            "M": 20,   # development_frame
            "N": 20,   # sovereignty_frame
            "O": 22,   # multilateral_frame
            "P": 26,   # responsibility_attribution
            "Q": 18,   # criticism_frame
            "R": 24,   # positive_self_image
            "S": 35,   # notes
        }

        for col, width in column_widths.items():
            worksheet.column_dimensions[col].width = width

        for row in worksheet.iter_rows():
            for cell in row:
                cell.alignment = Alignment(
                    vertical="top",
                    wrap_text=True
                )

        header_fill = PatternFill(
            fill_type="solid",
            fgColor="D9EAF7"
        )

        for cell in worksheet[1]:
            cell.font = Font(bold=True)
            cell.fill = header_fill

        warning_fill = PatternFill(
            fill_type="solid",
            fgColor="FFF2CC"
        )

        for row in range(2, worksheet.max_row + 1):
            needs_check_value = worksheet[f"J{row}"].value
            if needs_check_value == 1:
                for col in range(1, worksheet.max_column + 1):
                    worksheet.cell(row=row, column=col).fill = warning_fill


# =========================================================
# 7. 主程序
# =========================================================

def main():
    print("开始收集 2026年3月—4月 外交部例行记者会链接……")
    articles = collect_article_links()

    if not articles:
        print("没有找到目标文章。")
        print("可能原因：")
        print("1. 外交部官网分页结构变化；")
        print("2. 网络访问失败；")
        print("3. 时间范围内没有匹配标题。")
        return

    print(f"\n共找到 {len(articles)} 篇目标文章：")
    for item in articles:
        print(item["date"].strftime("%Y-%m-%d"), item["title"], item["url"])

    records = []

    print("\n开始抓取正文……")
    for idx, item in enumerate(articles, start=1):
        try:
            print(
                f"[{idx}/{len(articles)}] 抓取："
                f"{item['title']}（{item['date'].strftime('%Y-%m-%d')}）"
            )

            text = extract_article_text(item["url"])

            records.append({
                "date": item["date"],
                "date_str": item["date"].strftime("%Y-%m-%d"),
                "title": item["title"],
                "url": item["url"],
                "text": text,
            })

            time.sleep(0.7)

        except Exception as e:
            print(f"抓取失败：{item['url']}")
            print(f"错误信息：{e}")

    records.sort(key=lambda x: x["date"])

    print("\n开始保存文件……")
    save_txt(records)
    save_docx(records)
    save_excel(records)

    print("\n完成！已生成：")
    print(f"1. {OUT_TXT}")
    print(f"2. {OUT_DOCX}")
    print(f"3. {OUT_XLSX}")

    print("\n建议下一步：")
    print("打开 Excel 文件，筛选 needs_check = 1，检查未成功拆分或拆分异常的问答单元。")
    print("如果 needs_check 数量不多，可以手动修正；如果很多，再继续优化切分规则。")


if __name__ == "__main__":
    main()
